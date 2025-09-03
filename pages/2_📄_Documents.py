import streamlit as st
import requests
from bs4 import BeautifulSoup
import sqlite3
import re
import os
import tempfile
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd
from docx import Document
import openpyxl
from fpdf import FPDF
import base64
import json

# Load environment variables
load_dotenv()
client = OpenAI()
client.api_key = os.getenv('OPENAI_API_KEY')
OPENAI_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4o-mini')

st.set_page_config(page_title="Document Processing", layout="wide")

def init_documents_db():
    """Initialize documents database"""
    conn = sqlite3.connect('procurement.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            procurement_id TEXT,
            document_name TEXT,
            document_type TEXT,
            original_url TEXT,
            local_path TEXT,
            extracted_fields TEXT,
            form_data TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS form_submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER,
            form_data TEXT,
            pdf_path TEXT,
            submitted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (document_id) REFERENCES documents (id)
        )
    ''')
    
    conn.commit()
    conn.close()

def extract_document_fields_with_openai(document_content, document_type):
    """Extract form fields from document content using OpenAI LLM"""
    try:
        prompt = f"""
        You are an expert document analyzer specializing in Estonian procurement documents. 
        Analyze the following {document_type} document content and extract all form fields, requirements, and fillable sections.
        
        Document Content:
        {document_content[:8000]}  # Limit content to avoid token limits
        
        Please extract and return a JSON structure with the following information:
        {{
            "document_type": "type of document (contract, technical specification, etc.)",
            "title": "document title",
            "form_fields": [
                {{
                    "field_name": "name of the field",
                    "field_type": "text/number/date/dropdown/checkbox/textarea",
                    "label": "human readable label",
                    "required": true/false,
                    "description": "description of what this field is for",
                    "options": ["option1", "option2"] // only for dropdown fields
                }}
            ],
            "requirements": [
                "requirement 1",
                "requirement 2"
            ],
            "sections": [
                {{
                    "section_title": "section name",
                    "fields": ["field1", "field2"],
                    "description": "what this section covers"
                }}
            ],
            "key_information": {{
                "deadline": "if mentioned",
                "contact_person": "if mentioned",
                "submission_method": "how to submit",
                "evaluation_criteria": "how applications will be evaluated"
            }}
        }}
        
        Focus on identifying:
        1. All fillable fields and their types
        2. Required vs optional fields
        3. Validation requirements
        4. Section organization
        5. Key dates and contacts
        
        Return only valid JSON.
        """
        
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert document analyzer for Estonian procurement documents. Always return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        # Parse the JSON response
        extracted_data = json.loads(response.choices[0].message.content)
        return extracted_data
        
    except json.JSONDecodeError as e:
        st.error(f"Error parsing OpenAI response as JSON: {str(e)}")
        return None
    except Exception as e:
        st.error(f"Error extracting fields with OpenAI: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text content from DOCX file"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return None

def extract_text_from_xlsx(file_path):
    """Extract text content from XLSX file"""
    try:
        workbook = openpyxl.load_workbook(file_path)
        full_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            full_text.append(f"Sheet: {sheet_name}")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                if any(row_text):  # Only add non-empty rows
                    full_text.append('\t'.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error extracting text from XLSX: {str(e)}")
        return None

def simulate_document_fetch(procurement_id, document_name):
    """Simulate fetching a document from the procurement system"""
    # In production, this would actually download from the procurement system
    # For now, create sample document content
    
    sample_documents = {
        "Technical Specification": {
            "type": "docx",
            "content": """
            TECHNICAL SPECIFICATION FOR IT SERVICES
            
            1. GENERAL REQUIREMENTS
            Company Name: [TO BE FILLED]
            Contact Person: [TO BE FILLED]
            Email: [TO BE FILLED]
            Phone: [TO BE FILLED]
            
            2. TECHNICAL REQUIREMENTS
            Proposed Solution: [DESCRIBE YOUR SOLUTION]
            Technology Stack: [LIST TECHNOLOGIES]
            Implementation Timeline: [SPECIFY TIMELINE]
            Team Size: [NUMBER OF TEAM MEMBERS]
            
            3. EXPERIENCE
            Previous Projects: [LIST RELEVANT PROJECTS]
            Years of Experience: [NUMBER]
            Certifications: [LIST CERTIFICATIONS]
            
            4. PRICING
            Total Cost: [AMOUNT IN EUR]
            Payment Schedule: [DESCRIBE PAYMENT TERMS]
            
            5. COMPLIANCE
            ISO Certification: [YES/NO]
            GDPR Compliance: [YES/NO]
            Security Measures: [DESCRIBE]
            
            Signature: ________________
            Date: ________________
            """
        },
        "Contract Template": {
            "type": "docx",
            "content": """
            PROCUREMENT CONTRACT TEMPLATE
            
            CONTRACTING PARTIES:
            Procurer: [PROCURER NAME]
            Contractor: [CONTRACTOR NAME]
            Contract Number: [CONTRACT NUMBER]
            
            CONTRACT DETAILS:
            Service Description: [DESCRIBE SERVICES]
            Contract Value: [AMOUNT]
            Start Date: [DATE]
            End Date: [DATE]
            
            CONTRACTOR INFORMATION:
            Legal Name: [FULL LEGAL NAME]
            Registration Number: [REG NUMBER]
            Address: [FULL ADDRESS]
            Bank Account: [ACCOUNT NUMBER]
            
            PERFORMANCE REQUIREMENTS:
            Deliverables: [LIST DELIVERABLES]
            Quality Standards: [SPECIFY STANDARDS]
            Reporting Requirements: [SPECIFY REPORTING]
            
            TERMS AND CONDITIONS:
            Payment Terms: [PAYMENT SCHEDULE]
            Penalties: [PENALTY CLAUSES]
            Termination Conditions: [TERMINATION TERMS]
            
            Signatures:
            Procurer: ________________ Date: ________
            Contractor: ________________ Date: ________
            """
        },
        "Evaluation Criteria": {
            "type": "xlsx",
            "content": """
            EVALUATION CRITERIA FORM
            
            TECHNICAL EVALUATION (70 points):
            Technical Solution Quality: [SCORE 0-20]
            Team Qualifications: [SCORE 0-15]
            Implementation Plan: [SCORE 0-15]
            Previous Experience: [SCORE 0-20]
            
            FINANCIAL EVALUATION (30 points):
            Price Competitiveness: [SCORE 0-30]
            
            TOTAL SCORE: [SUM OF ALL SCORES]
            
            EVALUATOR INFORMATION:
            Evaluator Name: [NAME]
            Position: [POSITION]
            Date: [DATE]
            Signature: [SIGNATURE]
            """
        }
    }
    
    return sample_documents.get(document_name, {
        "type": "docx",
        "content": f"Sample content for {document_name}"
    })

def create_web_form(extracted_fields):
    """Create a web form based on extracted fields"""
    if not extracted_fields or 'form_fields' not in extracted_fields:
        st.error("No form fields found in the document")
        return None
    
    st.subheader("📝 Fill Out Form")
    st.write(f"**Document:** {extracted_fields.get('title', 'Unknown Document')}")
    
    form_data = {}
    
    with st.form("document_form"):
        # Group fields by sections if available
        sections = extracted_fields.get('sections', [])
        
        if sections:
            for section in sections:
                st.subheader(section['section_title'])
                if section.get('description'):
                    st.write(section['description'])
                
                # Find fields for this section
                section_fields = [field for field in extracted_fields['form_fields'] 
                                if field['field_name'] in section.get('fields', [])]
                
                for field in section_fields:
                    form_data[field['field_name']] = create_form_field(field)
        else:
            # No sections, display all fields
            for field in extracted_fields['form_fields']:
                form_data[field['field_name']] = create_form_field(field)
        
        # Submit button
        submitted = st.form_submit_button("💾 Save Form Data", type="primary")
        
        if submitted:
            return form_data
    
    return None

def create_form_field(field):
    """Create a Streamlit form field based on field definition"""
    field_name = field['field_name']
    field_type = field['field_type']
    label = field['label']
    required = field.get('required', False)
    description = field.get('description', '')
    
    # Add required indicator
    if required:
        label += " *"
    
    if field_type == 'text':
        return st.text_input(label, help=description, key=field_name)
    elif field_type == 'textarea':
        return st.text_area(label, help=description, key=field_name)
    elif field_type == 'number':
        return st.number_input(label, help=description, key=field_name)
    elif field_type == 'date':
        return st.date_input(label, help=description, key=field_name)
    elif field_type == 'dropdown':
        options = field.get('options', ['Option 1', 'Option 2'])
        return st.selectbox(label, options, help=description, key=field_name)
    elif field_type == 'checkbox':
        return st.checkbox(label, help=description, key=field_name)
    else:
        return st.text_input(label, help=description, key=field_name)

def generate_pdf_from_form_data(form_data, extracted_fields):
    """Generate PDF from filled form data"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Title
        title = extracted_fields.get('title', 'Procurement Form')
        pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
        pdf.ln(10)
        
        # Form data
        pdf.set_font('Arial', '', 12)
        
        for field_name, value in form_data.items():
            if value:  # Only include filled fields
                # Find field info
                field_info = next((f for f in extracted_fields['form_fields'] if f['field_name'] == field_name), None)
                label = field_info['label'] if field_info else field_name
                
                # Add field to PDF
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(0, 8, f"{label}:", ln=True)
                pdf.set_font('Arial', '', 10)
                
                # Handle different value types
                if isinstance(value, str):
                    # Split long text into multiple lines
                    if len(value) > 80:
                        lines = [value[i:i+80] for i in range(0, len(value), 80)]
                        for line in lines:
                            pdf.cell(0, 6, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                    else:
                        pdf.cell(0, 6, value.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                else:
                    pdf.cell(0, 6, str(value), ln=True)
                
                pdf.ln(2)
        
        # Save PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"form_submission_{timestamp}.pdf"
        pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
        pdf.output(pdf_path)
        
        return pdf_path
        
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None

def main():
    init_documents_db()
    
    st.title("📄 Document Processing & Form Generation")
    st.markdown("### AI-powered document analysis and form creation")
    
    # Document source selection
    col1, col2 = st.columns([2, 1])
    
    with col1:
        source_method = st.radio(
            "Document Source:",
            ["Fetch from Procurement", "Upload Document", "Test with Sample"],
            horizontal=True
        )
    
    with col2:
        st.info("💡 Choose how to get documents for processing")
    
    if source_method == "Fetch from Procurement":
        render_fetch_documents()
    elif source_method == "Upload Document":
        render_upload_documents()
    elif source_method == "Test with Sample":
        render_sample_documents()

def render_fetch_documents():
    st.subheader("📥 Fetch Documents from Procurement System")
    
    col1, col2 = st.columns(2)
    
    with col1:
        procurement_id = st.text_input("Procurement ID:", placeholder="e.g., 9262944")
    
    with col2:
        st.write("")  # Spacing
        fetch_button = st.button("🔍 Fetch Documents", type="primary")
    
    if fetch_button and procurement_id:
        with st.spinner("Fetching documents from procurement system..."):
            # Simulate fetching documents list
            sample_docs = [
                "Technical Specification",
                "Contract Template", 
                "Evaluation Criteria",
                "Terms and Conditions"
            ]
            
            st.success(f"Found {len(sample_docs)} documents for procurement {procurement_id}")
            
            # Display documents with fetch buttons
            for doc_name in sample_docs:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"📄 {doc_name}")
                
                with col2:
                    doc_type = "DOCX" if "Contract" in doc_name or "Specification" in doc_name else "PDF"
                    st.write(doc_type)
                
                with col3:
                    if st.button("📥 Process", key=f"fetch_{doc_name}"):
                        process_document(procurement_id, doc_name, "fetch")

def render_upload_documents():
    st.subheader("📤 Upload Document for Processing")
    
    uploaded_file = st.file_uploader(
        "Choose a document file",
        type=['docx', 'xlsx', 'pdf'],
        help="Upload Word documents (.docx), Excel files (.xlsx), or PDF files"
    )
    
    if uploaded_file is not None:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        st.success(f"Uploaded: {uploaded_file.name}")
        
        if st.button("🔍 Process Document", type="primary"):
            process_uploaded_document(tmp_file_path, uploaded_file.name)

def render_sample_documents():
    st.subheader("🧪 Test with Sample Documents")
    st.info("Use pre-loaded sample documents to test the AI extraction capabilities")
    
    sample_docs = [
        "Technical Specification",
        "Contract Template",
        "Evaluation Criteria"
    ]
    
    for doc_name in sample_docs:
        col1, col2, col3 = st.columns([3, 1, 1])
        
        with col1:
            st.write(f"📄 {doc_name}")
        
        with col2:
            st.write("Sample")
        
        with col3:
            if st.button("🧪 Test", key=f"sample_{doc_name}"):
                process_document("SAMPLE", doc_name, "sample")

def process_document(procurement_id, document_name, source_type):
    """Process a document and extract form fields"""
    st.markdown("---")
    st.subheader(f"🔍 Processing: {document_name}")
    
    with st.spinner("Analyzing document with AI..."):
        if source_type == "sample" or source_type == "fetch":
            # Get sample document content
            doc_data = simulate_document_fetch(procurement_id, document_name)
            document_content = doc_data['content']
            document_type = doc_data['type']
        else:
            st.error("Unsupported source type")
            return
        
        # Extract fields using OpenAI
        extracted_fields = extract_document_fields_with_openai(document_content, document_type)
        
        if extracted_fields:
            st.success("✅ Document analysis completed!")
            
            # Display extraction results
            display_extraction_results(extracted_fields)
            
            # Create web form
            form_data = create_web_form(extracted_fields)
            
            if form_data:
                st.success("Form data saved!")
                
                # Generate PDF option
                if st.button("📄 Generate PDF", type="primary"):
                    pdf_path = generate_pdf_from_form_data(form_data, extracted_fields)
                    
                    if pdf_path:
                        with open(pdf_path, "rb") as pdf_file:
                            pdf_bytes = pdf_file.read()
                        
                        st.download_button(
                            label="💾 Download Filled Form PDF",
                            data=pdf_bytes,
                            file_name=f"{document_name.replace(' ', '_')}_filled.pdf",
                            mime="application/pdf"
                        )
                        
                        st.success("PDF generated successfully!")
        else:
            st.error("Failed to extract fields from document")

def process_uploaded_document(file_path, filename):
    """Process an uploaded document"""
    st.markdown("---")
    st.subheader(f"🔍 Processing: {filename}")
    
    with st.spinner("Extracting text and analyzing with AI..."):
        # Extract text based on file type
        file_ext = filename.split('.')[-1].lower()
        
        if file_ext == 'docx':
            document_content = extract_text_from_docx(file_path)
        elif file_ext == 'xlsx':
            document_content = extract_text_from_xlsx(file_path)
        else:
            st.error("Unsupported file type. Please upload DOCX or XLSX files.")
            return
        
        if document_content:
            st.success("✅ Text extracted successfully!")
            
            # Show extracted text preview
            with st.expander("📄 Extracted Text Preview"):
                st.text_area("Document Content", document_content[:2000] + "..." if len(document_content) > 2000 else document_content, height=200)
            
            # Extract fields using OpenAI
            extracted_fields = extract_document_fields_with_openai(document_content, file_ext)
            
            if extracted_fields:
                st.success("✅ AI analysis completed!")
                
                # Display extraction results
                display_extraction_results(extracted_fields)
                
                # Create web form
                form_data = create_web_form(extracted_fields)
                
                if form_data:
                    st.success("Form data saved!")
                    
                    # Generate PDF option
                    if st.button("📄 Generate PDF", type="primary"):
                        pdf_path = generate_pdf_from_form_data(form_data, extracted_fields)
                        
                        if pdf_path:
                            with open(pdf_path, "rb") as pdf_file:
                                pdf_bytes = pdf_file.read()
                            
                            st.download_button(
                                label="💾 Download Filled Form PDF",
                                data=pdf_bytes,
                                file_name=f"{filename.split('.')[0]}_filled.pdf",
                                mime="application/pdf"
                            )
                            
                            st.success("PDF generated successfully!")
            else:
                st.error("Failed to analyze document with AI")
        else:
            st.error("Failed to extract text from document")

def display_extraction_results(extracted_fields):
    """Display the results of field extraction"""
    st.subheader("🤖 AI Extraction Results")
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📝 Form Fields", len(extracted_fields.get('form_fields', [])))
    
    with col2:
        required_fields = sum(1 for field in extracted_fields.get('form_fields', []) if field.get('required', False))
        st.metric("⚠️ Required Fields", required_fields)
    
    with col3:
        st.metric("📋 Sections", len(extracted_fields.get('sections', [])))
    
    with col4:
        st.metric("📄 Document Type", extracted_fields.get('document_type', 'Unknown'))
    
    # Detailed results in tabs
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Form Fields", "📋 Sections", "📋 Requirements", "ℹ️ Key Info"])
    
    with tab1:
        if extracted_fields.get('form_fields'):
            df_fields = pd.DataFrame(extracted_fields['form_fields'])
            st.dataframe(df_fields, use_container_width=True)
        else:
            st.info("No form fields detected")
    
    with tab2:
        if extracted_fields.get('sections'):
            for section in extracted_fields['sections']:
                st.write(f"**{section['section_title']}**")
                st.write(section.get('description', 'No description'))
                st.write(f"Fields: {', '.join(section.get('fields', []))}")
                st.divider()
        else:
            st.info("No sections detected")
    
    with tab3:
        if extracted_fields.get('requirements'):
            for i, req in enumerate(extracted_fields['requirements'], 1):
                st.write(f"{i}. {req}")
        else:
            st.info("No requirements detected")
    
    with tab4:
        key_info = extracted_fields.get('key_information', {})
        if key_info:
            for key, value in key_info.items():
                if value:
                    st.write(f"**{key.replace('_', ' ').title()}:** {value}")
        else:
            st.info("No key information detected")

if __name__ == "__main__":
    main()

