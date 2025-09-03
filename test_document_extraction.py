#!/usr/bin/env python3
"""
Test script for OpenAI LLM document field extraction
Tests the document processing capabilities and documents results
"""

import os
import json
import tempfile
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import openpyxl

# Load environment variables
load_dotenv()
client = OpenAI()
client.api_key = os.getenv('OPENAI_API_KEY')
OPENAI_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4.1-mini')

def create_sample_docx():
    """Create a sample DOCX document for testing"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Estonian Procurement Application Form', 0)
    
    # Add content
    doc.add_heading('1. Company Information', level=1)
    doc.add_paragraph('Company Name: ___________________________')
    doc.add_paragraph('Registration Number: ___________________________')
    doc.add_paragraph('Address: ___________________________')
    doc.add_paragraph('Contact Person: ___________________________')
    doc.add_paragraph('Email: ___________________________')
    doc.add_paragraph('Phone: ___________________________')
    
    doc.add_heading('2. Technical Proposal', level=1)
    doc.add_paragraph('Proposed Solution:')
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    
    doc.add_paragraph('Technology Stack: ___________________________')
    doc.add_paragraph('Implementation Timeline (months): ___________________________')
    doc.add_paragraph('Team Size: ___________________________')
    
    doc.add_heading('3. Experience and Qualifications', level=1)
    doc.add_paragraph('Years of Experience: ___________________________')
    doc.add_paragraph('Relevant Certifications: ___________________________')
    doc.add_paragraph('Previous Similar Projects (list 3):')
    doc.add_paragraph('1. ___________________________')
    doc.add_paragraph('2. ___________________________')
    doc.add_paragraph('3. ___________________________')
    
    doc.add_heading('4. Financial Proposal', level=1)
    doc.add_paragraph('Total Project Cost (EUR): ___________________________')
    doc.add_paragraph('Payment Schedule: ___________________________')
    doc.add_paragraph('Warranty Period (months): ___________________________')
    
    doc.add_heading('5. Compliance', level=1)
    doc.add_paragraph('ISO 27001 Certified: [ ] Yes [ ] No')
    doc.add_paragraph('GDPR Compliant: [ ] Yes [ ] No')
    doc.add_paragraph('Estonian Language Support: [ ] Yes [ ] No')
    
    doc.add_heading('6. Declaration', level=1)
    doc.add_paragraph('I hereby declare that all information provided is accurate and complete.')
    doc.add_paragraph('Signature: ___________________________')
    doc.add_paragraph('Date: ___________________________')
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def create_sample_xlsx():
    """Create a sample XLSX document for testing"""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Evaluation Form"
    
    # Headers
    sheet['A1'] = 'Estonian Procurement Evaluation Form'
    sheet['A1'].font = openpyxl.styles.Font(bold=True, size=14)
    
    # Evaluation criteria
    sheet['A3'] = 'Evaluation Criteria'
    sheet['B3'] = 'Max Points'
    sheet['C3'] = 'Awarded Points'
    sheet['D3'] = 'Comments'
    
    criteria = [
        ('Technical Solution Quality', 25),
        ('Team Qualifications', 20),
        ('Implementation Plan', 15),
        ('Previous Experience', 10),
        ('Price Competitiveness', 30)
    ]
    
    row = 4
    for criterion, max_points in criteria:
        sheet[f'A{row}'] = criterion
        sheet[f'B{row}'] = max_points
        sheet[f'C{row}'] = ''  # To be filled
        sheet[f'D{row}'] = ''  # To be filled
        row += 1
    
    # Total
    sheet[f'A{row+1}'] = 'TOTAL SCORE'
    sheet[f'B{row+1}'] = 100
    sheet[f'C{row+1}'] = f'=SUM(C4:C{row-1})'
    
    # Evaluator info
    sheet[f'A{row+3}'] = 'Evaluator Information'
    sheet[f'A{row+4}'] = 'Name:'
    sheet[f'A{row+5}'] = 'Position:'
    sheet[f'A{row+6}'] = 'Date:'
    sheet[f'A{row+7}'] = 'Signature:'
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    workbook.save(temp_file.name)
    return temp_file.name

def extract_text_from_docx(file_path):
    """Extract text content from DOCX file"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return None

def extract_text_from_xlsx(file_path):
    """Extract text content from XLSX file"""
    try:
        workbook = openpyxl.load_workbook(file_path)
        full_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            full_text.append(f"Sheet: {sheet_name}")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                if any(row_text):  # Only add non-empty rows
                    full_text.append('\t'.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from XLSX: {str(e)}")
        return None

def extract_document_fields_with_openai(document_content, document_type):
    """Extract form fields from document content using OpenAI LLM"""
    try:
        prompt = f"""
        You are an expert document analyzer specializing in Estonian procurement documents. 
        Analyze the following {document_type} document content and extract all form fields, requirements, and fillable sections.
        
        Document Content:
        {document_content[:8000]}  # Limit content to avoid token limits
        
        Please extract and return a JSON structure with the following information:
        {{
            "document_type": "type of document (contract, technical specification, etc.)",
            "title": "document title",
            "form_fields": [
                {{
                    "field_name": "name of the field",
                    "field_type": "text/number/date/dropdown/checkbox/textarea",
                    "label": "human readable label",
                    "required": true/false,
                    "description": "description of what this field is for",
                    "options": ["option1", "option2"] // only for dropdown fields
                }}
            ],
            "requirements": [
                "requirement 1",
                "requirement 2"
            ],
            "sections": [
                {{
                    "section_title": "section name",
                    "fields": ["field1", "field2"],
                    "description": "what this section covers"
                }}
            ],
            "key_information": {{
                "deadline": "if mentioned",
                "contact_person": "if mentioned",
                "submission_method": "how to submit",
                "evaluation_criteria": "how applications will be evaluated"
            }}
        }}
        
        Focus on identifying:
        1. All fillable fields and their types
        2. Required vs optional fields
        3. Validation requirements
        4. Section organization
        5. Key dates and contacts
        
        Return only valid JSON.
        """
        
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert document analyzer for Estonian procurement documents. Always return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        # Parse the JSON response
        extracted_data = json.loads(response.choices[0].message.content)
        return extracted_data
        
    except json.JSONDecodeError as e:
        print(f"Error parsing OpenAI response as JSON: {str(e)}")
        print(f"Raw response: {response.choices[0].message.content}")
        return None
    except Exception as e:
        print(f"Error extracting fields with OpenAI: {str(e)}")
        return None

def test_document_extraction():
    """Test the document extraction functionality"""
    print("🧪 Testing OpenAI LLM Document Field Extraction")
    print("=" * 60)
    
    # Test results
    results = {
        "timestamp": datetime.now().isoformat(),
        "model": OPENAI_MODEL,
        "tests": []
    }
    
    # Test 1: DOCX Document
    print("\n📄 Test 1: DOCX Document Extraction")
    print("-" * 40)
    
    docx_file = create_sample_docx()
    print(f"Created sample DOCX: {docx_file}")
    
    docx_content = extract_text_from_docx(docx_file)
    if docx_content:
        print(f"✅ Text extraction successful ({len(docx_content)} characters)")
        print(f"Preview: {docx_content[:200]}...")
        
        docx_fields = extract_document_fields_with_openai(docx_content, "docx")
        if docx_fields:
            print("✅ OpenAI field extraction successful")
            print(f"   - Document Type: {docx_fields.get('document_type', 'Unknown')}")
            print(f"   - Form Fields: {len(docx_fields.get('form_fields', []))}")
            print(f"   - Sections: {len(docx_fields.get('sections', []))}")
            print(f"   - Requirements: {len(docx_fields.get('requirements', []))}")
            
            results["tests"].append({
                "test_name": "DOCX Extraction",
                "status": "success",
                "document_type": docx_fields.get('document_type'),
                "fields_count": len(docx_fields.get('form_fields', [])),
                "sections_count": len(docx_fields.get('sections', [])),
                "requirements_count": len(docx_fields.get('requirements', [])),
                "extracted_data": docx_fields
            })
        else:
            print("❌ OpenAI field extraction failed")
            results["tests"].append({
                "test_name": "DOCX Extraction",
                "status": "failed",
                "error": "OpenAI extraction failed"
            })
    else:
        print("❌ Text extraction failed")
        results["tests"].append({
            "test_name": "DOCX Extraction",
            "status": "failed",
            "error": "Text extraction failed"
        })
    
    # Clean up
    os.unlink(docx_file)
    
    # Test 2: XLSX Document
    print("\n📊 Test 2: XLSX Document Extraction")
    print("-" * 40)
    
    xlsx_file = create_sample_xlsx()
    print(f"Created sample XLSX: {xlsx_file}")
    
    xlsx_content = extract_text_from_xlsx(xlsx_file)
    if xlsx_content:
        print(f"✅ Text extraction successful ({len(xlsx_content)} characters)")
        print(f"Preview: {xlsx_content[:200]}...")
        
        xlsx_fields = extract_document_fields_with_openai(xlsx_content, "xlsx")
        if xlsx_fields:
            print("✅ OpenAI field extraction successful")
            print(f"   - Document Type: {xlsx_fields.get('document_type', 'Unknown')}")
            print(f"   - Form Fields: {len(xlsx_fields.get('form_fields', []))}")
            print(f"   - Sections: {len(xlsx_fields.get('sections', []))}")
            print(f"   - Requirements: {len(xlsx_fields.get('requirements', []))}")
            
            results["tests"].append({
                "test_name": "XLSX Extraction",
                "status": "success",
                "document_type": xlsx_fields.get('document_type'),
                "fields_count": len(xlsx_fields.get('form_fields', [])),
                "sections_count": len(xlsx_fields.get('sections', [])),
                "requirements_count": len(xlsx_fields.get('requirements', [])),
                "extracted_data": xlsx_fields
            })
        else:
            print("❌ OpenAI field extraction failed")
            results["tests"].append({
                "test_name": "XLSX Extraction",
                "status": "failed",
                "error": "OpenAI extraction failed"
            })
    else:
        print("❌ Text extraction failed")
        results["tests"].append({
            "test_name": "XLSX Extraction",
            "status": "failed",
            "error": "Text extraction failed"
        })
    
    # Clean up
    os.unlink(xlsx_file)
    
    # Test 3: Complex Text Document
    print("\n📝 Test 3: Complex Text Document")
    print("-" * 40)
    
    complex_text = """
    ESTONIAN PUBLIC PROCUREMENT CONTRACT TEMPLATE
    
    SECTION A: CONTRACTOR INFORMATION
    Legal Company Name: [REQUIRED FIELD]
    Registration Number: [REQUIRED FIELD - Estonian format]
    VAT Number: [REQUIRED FIELD]
    Business Address: [REQUIRED FIELD]
    Postal Code: [REQUIRED FIELD - Estonian format]
    Contact Person: [REQUIRED FIELD]
    Email Address: [REQUIRED FIELD - valid email]
    Phone Number: [REQUIRED FIELD - Estonian format]
    
    SECTION B: PROJECT DETAILS
    Project Title: [REQUIRED FIELD]
    Project Description: [REQUIRED FIELD - minimum 500 characters]
    Estimated Start Date: [REQUIRED FIELD - date format DD.MM.YYYY]
    Estimated End Date: [REQUIRED FIELD - date format DD.MM.YYYY]
    Project Budget: [REQUIRED FIELD - EUR amount]
    
    SECTION C: TECHNICAL REQUIREMENTS
    Technology Stack: [REQUIRED FIELD - list technologies]
    Team Composition: [REQUIRED FIELD]
    Quality Assurance: [REQUIRED FIELD]
    Security Measures: [REQUIRED FIELD]
    
    SECTION D: COMPLIANCE
    ISO 27001 Certified: [CHECKBOX - Yes/No]
    GDPR Compliant: [CHECKBOX - Yes/No]
    Estonian Language Support: [CHECKBOX - Yes/No]
    EU Data Residency: [CHECKBOX - Yes/No]
    
    SECTION E: FINANCIAL TERMS
    Payment Schedule: [DROPDOWN - Monthly/Quarterly/Milestone-based]
    Currency: [DROPDOWN - EUR/USD]
    Warranty Period: [NUMBER - months]
    Penalty Clauses: [TEXTAREA]
    
    REQUIREMENTS:
    1. All contractors must be registered in Estonia or EU
    2. Minimum 3 years of relevant experience required
    3. Must provide 3 references from similar projects
    4. All documentation must be in Estonian or English
    5. Compliance with Estonian data protection laws mandatory
    
    EVALUATION CRITERIA:
    - Technical solution quality (40%)
    - Team qualifications (25%)
    - Price competitiveness (20%)
    - Previous experience (15%)
    
    SUBMISSION DEADLINE: 30 days from publication
    CONTACT: procurement@example.ee
    """
    
    complex_fields = extract_document_fields_with_openai(complex_text, "contract_template")
    if complex_fields:
        print("✅ Complex text extraction successful")
        print(f"   - Document Type: {complex_fields.get('document_type', 'Unknown')}")
        print(f"   - Form Fields: {len(complex_fields.get('form_fields', []))}")
        print(f"   - Sections: {len(complex_fields.get('sections', []))}")
        print(f"   - Requirements: {len(complex_fields.get('requirements', []))}")
        
        results["tests"].append({
            "test_name": "Complex Text Extraction",
            "status": "success",
            "document_type": complex_fields.get('document_type'),
            "fields_count": len(complex_fields.get('form_fields', [])),
            "sections_count": len(complex_fields.get('sections', [])),
            "requirements_count": len(complex_fields.get('requirements', [])),
            "extracted_data": complex_fields
        })
    else:
        print("❌ Complex text extraction failed")
        results["tests"].append({
            "test_name": "Complex Text Extraction",
            "status": "failed",
            "error": "OpenAI extraction failed"
        })
    
    # Save results
    results_file = f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\n📊 Test Results Summary")
    print("=" * 60)
    successful_tests = sum(1 for test in results["tests"] if test["status"] == "success")
    total_tests = len(results["tests"])
    print(f"✅ Successful Tests: {successful_tests}/{total_tests}")
    print(f"📄 Results saved to: {results_file}")
    
    if successful_tests == total_tests:
        print("🎉 All tests passed! OpenAI LLM extraction is working correctly.")
    else:
        print("⚠️ Some tests failed. Check the results file for details.")
    
    return results

if __name__ == "__main__":
    test_document_extraction()

