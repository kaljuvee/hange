#!/usr/bin/env python3
"""
Enhanced Document Processor for Hange AI
Implements all production recommendations:
1. Text preprocessing and caching
2. Enhanced Excel handling
3. Quality assurance with confidence scoring
4. User experience improvements
"""

import os
import json
import hashlib
import tempfile
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
import pandas as pd
import openpyxl
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import re
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
client = OpenAI()
client.api_key = os.getenv('OPENAI_API_KEY')
OPENAI_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4.1-mini')

@dataclass
class ExtractedField:
    """Represents an extracted form field with metadata"""
    field_name: str
    field_type: str
    label: str
    required: bool
    description: str
    options: Optional[List[str]] = None
    validation: Optional[str] = None
    confidence_score: float = 0.0
    source_text: Optional[str] = None
    needs_review: bool = False

@dataclass
class DocumentAnalysis:
    """Complete document analysis result"""
    document_type: str
    title: str
    form_fields: List[ExtractedField]
    requirements: List[str]
    sections: List[Dict[str, Any]]
    key_information: Dict[str, Any]
    confidence_score: float
    processing_time: float
    cache_hit: bool = False
    needs_human_review: bool = False

class DocumentCache:
    """Caching system for document extraction results"""
    
    def __init__(self, cache_db_path: str = "document_cache.db"):
        self.cache_db_path = cache_db_path
        self._init_cache_db()
    
    def _init_cache_db(self):
        """Initialize cache database"""
        conn = sqlite3.connect(self.cache_db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS document_cache (
                content_hash TEXT PRIMARY KEY,
                document_type TEXT,
                extraction_result TEXT,
                confidence_score REAL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                access_count INTEGER DEFAULT 1,
                last_accessed TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def get_content_hash(self, content: str) -> str:
        """Generate hash for document content"""
        return hashlib.sha256(content.encode()).hexdigest()
    
    def get_cached_result(self, content: str) -> Optional[Dict]:
        """Retrieve cached extraction result"""
        content_hash = self.get_content_hash(content)
        
        conn = sqlite3.connect(self.cache_db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT extraction_result, confidence_score 
            FROM document_cache 
            WHERE content_hash = ?
        ''', (content_hash,))
        
        result = cursor.fetchone()
        
        if result:
            # Update access statistics
            cursor.execute('''
                UPDATE document_cache 
                SET access_count = access_count + 1, 
                    last_accessed = CURRENT_TIMESTAMP 
                WHERE content_hash = ?
            ''', (content_hash,))
            conn.commit()
            
            extraction_result = json.loads(result[0])
            extraction_result['cache_hit'] = True
            logger.info(f"Cache hit for document hash: {content_hash[:8]}...")
            
        conn.close()
        return extraction_result if result else None
    
    def cache_result(self, content: str, document_type: str, extraction_result: Dict):
        """Cache extraction result"""
        content_hash = self.get_content_hash(content)
        
        conn = sqlite3.connect(self.cache_db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT OR REPLACE INTO document_cache 
            (content_hash, document_type, extraction_result, confidence_score)
            VALUES (?, ?, ?, ?)
        ''', (content_hash, document_type, json.dumps(extraction_result), 
              extraction_result.get('confidence_score', 0.0)))
        
        conn.commit()
        conn.close()
        logger.info(f"Cached result for document hash: {content_hash[:8]}...")

class EnhancedDocumentProcessor:
    """Enhanced document processor with production features"""
    
    def __init__(self):
        self.cache = DocumentCache()
        self.confidence_threshold = 0.7
        self.review_threshold = 0.5
    
    def preprocess_text(self, text: str) -> str:
        """Enhanced text preprocessing"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize Estonian characters
        text = text.replace('õ', 'õ').replace('ä', 'ä').replace('ö', 'ö').replace('ü', 'ü')
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Normalize currency symbols
        text = re.sub(r'€|EUR|euro', 'EUR', text, flags=re.IGNORECASE)
        
        # Normalize phone numbers
        text = re.sub(r'\+372\s*(\d{3,4})\s*(\d{4})', r'+372 \1 \2', text)
        
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Enhanced DOCX text extraction"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(f"HEADER: {paragraph.text.strip()}")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(f"FOOTER: {paragraph.text.strip()}")
            
            return '\n'.join(full_text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return None
    
    def extract_text_from_xlsx(self, file_path: str) -> str:
        """Enhanced XLSX text extraction with structure analysis"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            full_text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                full_text.append(f"SHEET: {sheet_name}")
                
                # Analyze sheet structure
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                # Extract headers (first few rows)
                headers = []
                for row in range(1, min(4, max_row + 1)):
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell = sheet.cell(row=row, column=col)
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data:
                        headers.append(' | '.join(row_data))
                
                if headers:
                    full_text.append("HEADERS:")
                    full_text.extend(headers)
                
                # Extract data rows (sample)
                data_rows = []
                for row in range(4, min(max_row + 1, 20)):  # Sample first 16 data rows
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell = sheet.cell(row=row, column=col)
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data and any(row_data):  # Only non-empty rows
                        data_rows.append(' | '.join(row_data))
                
                if data_rows:
                    full_text.append("DATA SAMPLE:")
                    full_text.extend(data_rows[:5])  # First 5 data rows
                
                # Extract formulas and validation
                formulas = []
                for row in sheet.iter_rows():
                    for cell in row:
                        if hasattr(cell, 'data_type') and cell.data_type == 'f':
                            formulas.append(f"FORMULA in {cell.coordinate}: {cell.value}")
                        
                        # Check for data validation
                        if cell.data_validation and cell.data_validation.formula1:
                            full_text.append(f"VALIDATION in {cell.coordinate}: {cell.data_validation.formula1}")
                
                if formulas:
                    full_text.append("FORMULAS:")
                    full_text.extend(formulas[:5])  # First 5 formulas
            
            return '\n'.join(full_text)
            
        except Exception as e:
            logger.error(f"Error extracting text from XLSX: {str(e)}")
            return None
    
    def calculate_confidence_score(self, extracted_data: Dict) -> float:
        """Calculate confidence score for extraction"""
        score = 0.0
        total_weight = 0.0
        
        # Field completeness (30%)
        fields = extracted_data.get('form_fields', [])
        if fields:
            complete_fields = sum(1 for field in fields if field.get('field_name') and field.get('field_type'))
            score += (complete_fields / len(fields)) * 0.3
        total_weight += 0.3
        
        # Field type accuracy (25%)
        valid_types = ['text', 'number', 'email', 'tel', 'date', 'checkbox', 'dropdown', 'textarea']
        if fields:
            valid_type_fields = sum(1 for field in fields if field.get('field_type') in valid_types)
            score += (valid_type_fields / len(fields)) * 0.25
        total_weight += 0.25
        
        # Requirements extraction (20%)
        requirements = extracted_data.get('requirements', [])
        if requirements:
            score += min(len(requirements) / 5, 1.0) * 0.2  # Up to 5 requirements
        total_weight += 0.2
        
        # Sections organization (15%)
        sections = extracted_data.get('sections', [])
        if sections:
            score += min(len(sections) / 4, 1.0) * 0.15  # Up to 4 sections
        total_weight += 0.15
        
        # Key information extraction (10%)
        key_info = extracted_data.get('key_information', {})
        if key_info:
            non_null_keys = sum(1 for v in key_info.values() if v is not None)
            score += min(non_null_keys / 4, 1.0) * 0.1  # Up to 4 key info items
        total_weight += 0.1
        
        return score / total_weight if total_weight > 0 else 0.0
    
    def extract_fields_with_fallback(self, content: str, document_type: str) -> Dict:
        """Extract fields with fallback mechanisms"""
        try:
            # Primary extraction with OpenAI
            result = self._extract_with_openai(content, document_type)
            
            if result and self.calculate_confidence_score(result) >= self.review_threshold:
                return result
            
        except Exception as e:
            logger.warning(f"OpenAI extraction failed: {str(e)}")
        
        # Fallback to rule-based extraction
        logger.info("Using fallback rule-based extraction")
        return self._extract_with_rules(content, document_type)
    
    def _extract_with_openai(self, content: str, document_type: str) -> Dict:
        """Extract fields using OpenAI with enhanced prompt"""
        prompt = f"""
        You are an expert document analyzer specializing in Estonian procurement documents. 
        Analyze the following {document_type} document and extract all form fields with high accuracy.
        
        Document Content:
        {content[:12000]}  # Increased content limit
        
        Return a JSON structure with confidence scoring:
        {{
            "document_type": "specific document type",
            "title": "document title",
            "form_fields": [
                {{
                    "field_name": "snake_case_name",
                    "field_type": "text|number|email|tel|date|checkbox|dropdown|textarea",
                    "label": "human readable Estonian label",
                    "required": true/false,
                    "description": "detailed field description",
                    "options": ["option1", "option2"], // only for dropdown/checkbox
                    "validation": "validation rules (e.g., min_length:5, pattern:^[0-9]+$)",
                    "confidence_score": 0.0-1.0,
                    "source_text": "original text that led to this field"
                }}
            ],
            "requirements": ["requirement 1", "requirement 2"],
            "sections": [
                {{
                    "section_title": "section name",
                    "fields": ["field1", "field2"],
                    "description": "section purpose"
                }}
            ],
            "key_information": {{
                "deadline": "if mentioned",
                "contact_person": "if mentioned",
                "submission_method": "how to submit",
                "evaluation_criteria": "evaluation details"
            }}
        }}
        
        Focus on:
        1. Estonian language field labels and descriptions
        2. Proper validation rules for Estonian formats (phone, ID, VAT)
        3. High confidence scores for clearly identifiable fields
        4. Detailed source text references
        
        Return only valid JSON.
        """
        
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert Estonian procurement document analyzer. Always return valid JSON with confidence scores."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        return json.loads(response.choices[0].message.content)
    
    def _extract_with_rules(self, content: str, document_type: str) -> Dict:
        """Fallback rule-based extraction"""
        fields = []
        requirements = []
        
        # Common Estonian procurement field patterns
        field_patterns = {
            'company_name': r'(ettevõtte?\s+nimi|firma\s+nimi|ärinimi)',
            'registration_number': r'(registrikood|reg\.?\s*kood)',
            'vat_number': r'(kmkr?\s+number|käibemaksu)',
            'contact_person': r'(kontaktisik|vastutav\s+isik)',
            'email': r'(e-?post|email)',
            'phone': r'(telefon|tel\.?)',
            'address': r'(aadress|asukoht)',
        }
        
        for field_name, pattern in field_patterns.items():
            if re.search(pattern, content, re.IGNORECASE):
                fields.append({
                    'field_name': field_name,
                    'field_type': 'email' if 'email' in field_name else 'text',
                    'label': field_name.replace('_', ' ').title(),
                    'required': True,
                    'description': f'Auto-detected {field_name}',
                    'confidence_score': 0.6,
                    'source_text': 'Rule-based detection'
                })
        
        # Extract requirements
        req_patterns = [
            r'nõue[d]?:?\s*(.+)',
            r'tingimus[ed]?:?\s*(.+)',
            r'kriteerium[id]?:?\s*(.+)'
        ]
        
        for pattern in req_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches[:3])  # Max 3 requirements
        
        return {
            'document_type': document_type,
            'title': 'Rule-based extraction',
            'form_fields': fields,
            'requirements': requirements,
            'sections': [],
            'key_information': {}
        }
    
    def process_document(self, file_path: str, document_type: str = None) -> DocumentAnalysis:
        """Main document processing method with all enhancements"""
        start_time = datetime.now()
        
        # Determine document type
        if not document_type:
            ext = os.path.splitext(file_path)[1].lower()
            document_type = ext[1:] if ext else 'unknown'
        
        # Extract text based on file type
        if document_type == 'docx':
            raw_text = self.extract_text_from_docx(file_path)
        elif document_type == 'xlsx':
            raw_text = self.extract_text_from_xlsx(file_path)
        else:
            # Try reading as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
            except:
                raise ValueError(f"Unsupported file type: {document_type}")
        
        if not raw_text:
            raise ValueError("Could not extract text from document")
        
        # Preprocess text
        processed_text = self.preprocess_text(raw_text)
        
        # Check cache first
        cached_result = self.cache.get_cached_result(processed_text)
        if cached_result:
            processing_time = (datetime.now() - start_time).total_seconds()
            # Remove cache_hit from cached_result to avoid conflict
            cached_result_clean = {k: v for k, v in cached_result.items() if k != 'cache_hit'}
            return DocumentAnalysis(
                **cached_result_clean,
                processing_time=processing_time,
                cache_hit=True
            )
        
        # Extract fields with fallback
        extraction_result = self.extract_fields_with_fallback(processed_text, document_type)
        
        # Calculate confidence score
        confidence_score = self.calculate_confidence_score(extraction_result)
        extraction_result['confidence_score'] = confidence_score
        
        # Determine if human review is needed
        needs_review = confidence_score < self.confidence_threshold
        extraction_result['needs_human_review'] = needs_review
        
        # Cache the result
        self.cache.cache_result(processed_text, document_type, extraction_result)
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Convert to DocumentAnalysis
        form_fields = [ExtractedField(**field) for field in extraction_result.get('form_fields', [])]
        
        return DocumentAnalysis(
            document_type=extraction_result.get('document_type', document_type),
            title=extraction_result.get('title', 'Untitled'),
            form_fields=form_fields,
            requirements=extraction_result.get('requirements', []),
            sections=extraction_result.get('sections', []),
            key_information=extraction_result.get('key_information', {}),
            confidence_score=confidence_score,
            processing_time=processing_time,
            cache_hit=False,
            needs_human_review=needs_review
        )
    
    def validate_extracted_fields(self, fields: List[ExtractedField]) -> List[Dict[str, str]]:
        """Validate extracted fields and return validation errors"""
        errors = []
        
        for field in fields:
            # Check required fields
            if field.required and not field.field_name:
                errors.append({
                    'field': field.label,
                    'error': 'Required field name is missing'
                })
            
            # Validate field types
            valid_types = ['text', 'number', 'email', 'tel', 'date', 'checkbox', 'dropdown', 'textarea']
            if field.field_type not in valid_types:
                errors.append({
                    'field': field.label,
                    'error': f'Invalid field type: {field.field_type}'
                })
            
            # Validate Estonian phone pattern
            if field.field_type == 'tel' and field.validation:
                if 'pattern:' in field.validation and '+372' not in field.validation:
                    errors.append({
                        'field': field.label,
                        'error': 'Phone validation should include Estonian format (+372)'
                    })
            
            # Check confidence score
            if field.confidence_score < self.review_threshold:
                errors.append({
                    'field': field.label,
                    'error': f'Low confidence score: {field.confidence_score:.2f}'
                })
        
        return errors
    
    def generate_form_preview(self, analysis: DocumentAnalysis) -> str:
        """Generate HTML form preview"""
        html = f"""
        <div class="form-preview">
            <h2>{analysis.title}</h2>
            <p><strong>Document Type:</strong> {analysis.document_type}</p>
            <p><strong>Confidence Score:</strong> {analysis.confidence_score:.2%}</p>
            <p><strong>Processing Time:</strong> {analysis.processing_time:.2f}s</p>
            
            {'<div class="alert alert-warning">⚠️ This document needs human review</div>' if analysis.needs_human_review else ''}
            
            <form>
        """
        
        for field in analysis.form_fields:
            required_attr = 'required' if field.required else ''
            confidence_class = 'low-confidence' if field.confidence_score < self.confidence_threshold else ''
            
            html += f"""
                <div class="form-group {confidence_class}">
                    <label for="{field.field_name}">
                        {field.label} {'*' if field.required else ''}
                        <span class="confidence">({field.confidence_score:.1%})</span>
                    </label>
            """
            
            if field.field_type == 'textarea':
                html += f'<textarea id="{field.field_name}" name="{field.field_name}" {required_attr} placeholder="{field.description}"></textarea>'
            elif field.field_type == 'dropdown':
                html += f'<select id="{field.field_name}" name="{field.field_name}" {required_attr}>'
                if field.options:
                    for option in field.options:
                        html += f'<option value="{option}">{option}</option>'
                html += '</select>'
            elif field.field_type == 'checkbox':
                html += f'<input type="checkbox" id="{field.field_name}" name="{field.field_name}" {required_attr}>'
            else:
                html += f'<input type="{field.field_type}" id="{field.field_name}" name="{field.field_name}" {required_attr} placeholder="{field.description}">'
            
            html += f'<small class="help-text">{field.description}</small>'
            html += '</div>'
        
        html += """
            </form>
            
            <div class="requirements">
                <h3>Requirements</h3>
                <ul>
        """
        
        for req in analysis.requirements:
            html += f'<li>{req}</li>'
        
        html += """
                </ul>
            </div>
        </div>
        """
        
        return html

# Example usage and testing
if __name__ == "__main__":
    processor = EnhancedDocumentProcessor()
    
    # Test with sample data
    print("Enhanced Document Processor initialized successfully!")
    print(f"Cache database: {processor.cache.cache_db_path}")
    print(f"Confidence threshold: {processor.confidence_threshold}")
    print(f"Review threshold: {processor.review_threshold}")

